"""Generate the Arabic Word report for Al-Farouk's first 5 SOC-upgrade tasks."""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "تقرير_مهام_الفاروق.docx")

ARABIC_FONT = "Calibri"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)


def set_run_arabic(run, size=12, bold=False, color=None):
    run.font.name = ARABIC_FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), ARABIC_FONT)
    rPr_lang = rPr.makeelement(qn("w:lang"), {qn("w:bidi"): "ar-IQ"})
    rPr.append(rPr_lang)


def add_par(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, color=None, space_after=6):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_arabic(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 15, 3: 13}
    colors = {1: (0x1F, 0x4E, 0x79), 2: (0x2E, 0x75, 0xB6), 3: (0x40, 0x40, 0x40)}
    p = add_par(doc, text, size=sizes.get(level, 13), bold=True, color=colors.get(level, (0, 0, 0)), space_after=8)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_arabic(run, size=size)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return p


def main():
    doc = Document()

    # Default document direction / font
    style = doc.styles["Normal"]
    style.font.name = ARABIC_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), ARABIC_FONT)

    # Title
    add_par(doc, "تقرير إنجاز المهام", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1F, 0x4E, 0x79), space_after=2)
    add_par(doc, "مهندس الفاروق – الأنظمة والاستجابة (Systems & Response)", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x40, 0x40, 0x40), space_after=2)
    add_par(doc, "مشروع: نظام الدفاع السيبراني (Cyber Defense System) – خطة ترقية SOC", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x60, 0x60, 0x60), space_after=2)
    add_par(doc, "التاريخ: 2026-06-10", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x60, 0x60, 0x60), space_after=12)

    # Intro
    add_heading(doc, "مقدمة", level=1)
    add_par(doc,
            "بناءً على تقسيم المهام الوارد في خطة ترقية النظام إلى منصة SOC (soc_upgrade_plan_ar.pdf)، "
            "تم إنجاز المهام الخمس الأولى المسندة إلى مهندس الأنظمة والاستجابة (الفاروق)، "
            "وجميعها تندرج ضمن البند الثالث من الخطة: «الميزات الأمنية السيبرانية المتقدمة» (الصفحة 4، القسمان أ و ب). "
            "تم تنفيذ جميع المهام مع الحفاظ التام على نسق المشروع الحالي: نفس بنية الطبقات "
            "(storage / backend / routes)، نفس آلية الصلاحيات RBAC، ونفس أسلوب كتابة الاختبارات "
            "(pytest + FastAPI TestClient)، بحيث تندمج الإضافات بسلاسة مع بقية فريق العمل.")

    add_heading(doc, "ملخص المهام المنجزة", level=2)
    for t in [
        "1. قاعدة بيانات الحوادث الأمنية (Incident Database Schema)",
        "2. إدارة الحالات (Case Management)",
        "3. سجلات تدقيق غير قابلة للتعديل بتقنية تسلسل التجزئة (Immutable Audit Logs – Hash Chaining)",
        "4. محرك إعادة تشغيل الأحداث (Event Replay Engine)",
        "5. محرك الاستعلامات المتقدم (Advanced Query Engine)",
    ]:
        add_bullet(doc, t)

    # ------------------------------------------------------------------
    # Task 1
    # ------------------------------------------------------------------
    add_heading(doc, "المهمة 1: قاعدة بيانات الحوادث الأمنية (Incident Database Schema)", level=1)
    add_par(doc, "الوصف:", bold=True)
    add_par(doc,
            "تمت إضافة بنية بيانات كاملة لإدارة «الحوادث الأمنية» (Security Incidents) كوحدة منفصلة عن "
            "الأحداث الخام (security_events)، بحيث يمكن لفريق الاستجابة تجميع عدة أحداث ضمن حادثة واحدة، "
            "ومتابعة دورة حياتها من الاكتشاف وحتى الإغلاق.")

    add_par(doc, "التغييرات التقنية:", bold=True)
    add_bullet(doc, "إنشاء جدول security_incidents يحتوي: incident_id (مفتاح فريد)، title، description، "
                    "severity، status، assigned_to، case_id، created_at، updated_at، notes.")
    add_bullet(doc, "إنشاء جدول incident_events لربط الأحداث (security_events) بالحوادث (علاقة واحد-لعدة).")
    add_bullet(doc, "حالات الحادثة (Incident Lifecycle) محصورة ضمن: open → investigating → pending_approval → resolved، "
                    "مع رفض أي حالة غير معروفة عبر استثناء ValueError.")
    add_bullet(doc, "تمت إضافة هذه البنية في كل من storage/persistence.py (SQLite) و storage/postgres_store.py "
                    "(PostgreSQL) بنفس الأسلوب المتّبع في الجداول الحالية، ضماناً لعمل النظام على الخيارين.")

    add_par(doc, "الدوال المضافة في طبقة التخزين:", bold=True)
    add_code(doc, "create_incident()    -> إنشاء حادثة جديدة وربطها بالأحداث الأولية\n"
                  "get_incident()       -> جلب حادثة واحدة مع الأحداث المرتبطة بها\n"
                  "list_incidents()     -> قائمة الحوادث مع فلاتر (status, severity, assigned_to, case_id)\n"
                  "update_incident()    -> تعديل بيانات/حالة الحادثة (مع التحقق من صحة الحالة)\n"
                  "link_event_to_incident() / list_incident_events()")

    add_par(doc, "واجهات برمجة التطبيقات (API) في backend/routes/incidents.py:", bold=True)
    add_code(doc, "GET    /api/incidents              قائمة الحوادث (مع فلاتر)\n"
                  "POST   /api/incidents              إنشاء حادثة جديدة\n"
                  "GET    /api/incidents/{id}         تفاصيل حادثة + الأحداث المرتبطة\n"
                  "PATCH  /api/incidents/{id}         تعديل بيانات/حالة الحادثة\n"
                  "POST   /api/incidents/{id}/events  ربط حدث أمني إضافي بالحادثة")

    add_par(doc, "الصلاحيات (RBAC):", bold=True)
    add_par(doc, "تمت إضافة صلاحيتين جديدتين في security/roles.py: incidents:read (متاحة لـ admin، analyst، viewer) "
                  "و incidents:write (متاحة لـ admin و analyst فقط)، بما يطابق نمط الصلاحيات الحالي للنظام.")

    # ------------------------------------------------------------------
    # Task 2
    # ------------------------------------------------------------------
    add_heading(doc, "المهمة 2: إدارة الحالات (Case Management)", level=1)
    add_par(doc, "الوصف:", bold=True)
    add_par(doc,
            "أُضيفت طبقة «الحالة» (Case) فوق طبقة الحوادث، لتمكين فريق الاستجابة من تجميع عدة حوادث "
            "مرتبطة ببعضها (مثل حملة هجوم واحدة متعددة المراحل) ضمن حالة تحقيق واحدة، ومتابعتها ككيان واحد.")

    add_par(doc, "التغييرات التقنية:", bold=True)
    add_bullet(doc, "إنشاء جدول security_cases يحتوي: case_id، title، description، status، created_at، "
                    "updated_at، notes.")
    add_bullet(doc, "إنشاء جدول ربط case_incidents (علاقة عديد-لعديد) بين الحالات والحوادث.")
    add_bullet(doc, "عند ربط حادثة بحالة، يتم تحديث الحقل case_id في جدول الحوادث تلقائياً، لتسهيل الاستعلام "
                    "المباشر عن «إلى أي حالة تنتمي هذه الحادثة».")
    add_bullet(doc, "حالات الحالة (Case Status) محصورة ضمن: open → investigating → closed.")

    add_par(doc, "الدوال المضافة في طبقة التخزين:", bold=True)
    add_code(doc, "create_case()            -> إنشاء حالة جديدة وربطها بحوادث أولية\n"
                  "get_case()               -> جلب حالة مع قائمة الحوادث المرتبطة بها\n"
                  "list_cases()             -> قائمة الحالات مع فلتر الحالة (status)\n"
                  "update_case()            -> تعديل بيانات/حالة الحالة\n"
                  "link_incident_to_case()  -> ربط حادثة إضافية بحالة قائمة")

    add_par(doc, "واجهات برمجة التطبيقات (API):", bold=True)
    add_code(doc, "GET    /api/cases                  قائمة الحالات\n"
                  "POST   /api/cases                  إنشاء حالة جديدة (مع حوادث أولية اختيارية)\n"
                  "GET    /api/cases/{id}              تفاصيل الحالة + الحوادث المرتبطة\n"
                  "PATCH  /api/cases/{id}              تعديل بيانات/حالة الحالة\n"
                  "POST   /api/cases/{id}/incidents    ربط حادثة إضافية بالحالة")

    add_par(doc, "الصلاحيات (RBAC):", bold=True)
    add_par(doc, "تمت إضافة صلاحيتين: cases:read (admin، analyst، viewer) و cases:write (admin، analyst).")

    # ------------------------------------------------------------------
    # Task 3
    # ------------------------------------------------------------------
    add_heading(doc, "المهمة 3: سجلات التدقيق غير القابلة للتعديل (Immutable Audit Logs – Hash Chaining)", level=1)
    add_par(doc, "الوصف:", bold=True)
    add_par(doc,
            "تم تطوير نظام سجلات تدقيق (Audit Trail) يسجّل كل إجراء حساس يقوم به المستخدمون "
            "(إنشاء/تعديل حادثة أو حالة، حظر/فك حظر عنوان IP، ...إلخ)، مع ضمان عدم إمكانية التلاعب بالسجلات "
            "بأثر رجعي دون كشف ذلك، باستخدام تقنية تسلسل التجزئة (Hash Chaining) على غرار سلاسل الكتل (Blockchain).")

    add_par(doc, "آلية العمل:", bold=True)
    add_bullet(doc, "كل سجل تدقيق يحتوي على: log_id (تسلسلي تلقائي)، timestamp، actor (المستخدم المنفِّذ)، "
                    "action (نوع الإجراء)، details (تفاصيل الإجراء بصيغة JSON)، prev_hash (تجزئة السجل السابق)، "
                    "hash (تجزئة هذا السجل).")
    add_bullet(doc, "أول سجل في السلسلة يستخدم قيمة ابتدائية ثابتة GENESIS_HASH (64 صفراً)، تماماً كما في كتلة "
                    "التكوين (Genesis Block) في سلاسل الكتل.")
    add_bullet(doc, "يتم حساب hash لكل سجل عبر دالة compute_audit_hash() باستخدام خوارزمية SHA-256 على بيانات "
                    "السجل مرتّبة بصيغة JSON قياسية (sorted keys)، بحيث أي تعديل ولو بحرف واحد في details "
                    "يغيّر قيمة hash بالكامل.")
    add_bullet(doc, "كل سجل جديد يأخذ hash السجل السابق كقيمة prev_hash له، مما يكوّن «سلسلة» متصلة: أي تعديل "
                    "على سجل قديم يكسر تطابق hash مع prev_hash الخاص بالسجل الذي يليه.")

    add_par(doc, "دالة التحقق من سلامة السلسلة:", bold=True)
    add_par(doc,
            "verify_audit_chain() تعيد إعادة حساب التجزئة لكل سجل بالترتيب والتحقق من تطابقها مع المخزّن، "
            "وتعيد كائن نتيجة بالشكل: { valid: true/false, total: عدد السجلات, broken_at: رقم أول سجل تم العبث به أو null }.")

    add_par(doc, "واجهات برمجة التطبيقات (backend/routes/forensics.py):", bold=True)
    add_code(doc, "GET /api/audit-logs           عرض جميع سجلات التدقيق (الأحدث أولاً)\n"
                  "GET /api/audit-logs/verify    التحقق من سلامة سلسلة السجلات بالكامل")

    add_par(doc, "الربط مع الإجراءات الحالية في النظام:", bold=True)
    add_par(doc,
            "تم تفعيل تسجيل التدقيق تلقائياً عند: إنشاء/تعديل حادثة، إنشاء/تعديل حالة، ربط حدث أو حادثة، "
            "وكذلك عند حظر أو فك حظر عنوان IP في backend/routes/security.py (الإجراءات الموجودة مسبقاً في "
            "النظام)، بحيث أصبح كل إجراء استجابة موثّقاً في سجل تدقيق غير قابل للتزوير.")

    add_par(doc, "الصلاحيات (RBAC):", bold=True)
    add_par(doc, "تمت إضافة صلاحية audit:verify (متاحة لـ admin و analyst) للتحكم بالوصول إلى سجلات التدقيق.")

    # ------------------------------------------------------------------
    # Task 4
    # ------------------------------------------------------------------
    add_heading(doc, "المهمة 4: محرك إعادة تشغيل الأحداث (Event Replay Engine)", level=1)
    add_par(doc, "الوصف:", bold=True)
    add_par(doc,
            "تم تطوير محرك يتيح لفريق التحقيق الجنائي الرقمي (Digital Forensics) استعراض جميع الأحداث "
            "الأمنية المرتبطة بكيان معيّن (جهاز، مستخدم، عنوان IP، الجدار الناري...) بترتيب زمني تصاعدي "
            "(من الأقدم إلى الأحدث)، بهدف «إعادة تشغيل» تسلسل الهجوم خطوة بخطوة كما حدث فعلياً.")

    add_par(doc, "آلية العمل:", bold=True)
    add_bullet(doc, "الدالة replay_events(target_entity, start_ts, end_ts, limit) في طبقة التخزين تُرجع "
                    "الأحداث المرتبطة بالكيان target_entity، مرتبة تصاعدياً حسب timestamp (وليس تنازلياً "
                    "كما هو معتاد في عرض الأحداث الحديثة).")
    add_bullet(doc, "يمكن تحديد نافذة زمنية اختيارية (start_ts و end_ts) لتضييق نطاق إعادة التشغيل على "
                    "فترة زمنية محددة من الحادثة.")

    add_par(doc, "واجهة برمجة التطبيقات:", bold=True)
    add_code(doc, "GET /api/forensics/replay?target_entity=...&start_ts=...&end_ts=...&limit=...")

    add_par(doc, "الصلاحيات (RBAC):", bold=True)
    add_par(doc, "تمت إضافة صلاحية forensics:read (متاحة لـ admin و analyst فقط، وليس viewer)، حماية "
                  "لخصوصية بيانات التحقيق الجنائي.")

    # ------------------------------------------------------------------
    # Task 5
    # ------------------------------------------------------------------
    add_heading(doc, "المهمة 5: محرك الاستعلامات المتقدم (Advanced Query Engine)", level=1)
    add_par(doc, "الوصف:", bold=True)
    add_par(doc,
            "تم تطوير محرك استعلامات مرن يتيح لفريق التحليل البحث ضمن الأحداث الأمنية باستخدام عدة شروط "
            "مركّبة (Compound Filters) مع منطق AND/OR، دون الحاجة لكتابة استعلامات SQL مباشرة، ومع ضمان "
            "أمان كامل ضد هجمات حقن SQL (SQL Injection).")

    add_par(doc, "آلية العمل والأمان:", bold=True)
    add_bullet(doc, "تم تعريف قائمة بيضاء (Whitelist) للحقول المسموح بالاستعلام عليها QUERY_FIELDS "
                    "(event_id، timestamp، event_type، severity، source_ip، target_entity، description، "
                    "action_taken، status) وقائمة بيضاء للعمليات QUERY_OPERATORS "
                    "(eq، neq، gt، gte، lt، lte، contains، in).")
    add_bullet(doc, "أي حقل أو عملية غير موجودة في القائمتين يتم تجاهلها بهدوء دون التسبب بخطأ أو ثغرة أمنية "
                    "(تم اختبار ذلك صراحةً).")
    add_bullet(doc, "دالة build_event_query() تبني جملة SQL مع شروط (placeholders) معاملة بأمان، وتدعم كلا "
                    "من SQLite (?) و PostgreSQL (%s) عبر معامل placeholder، بما يتوافق مع طبقة التخزين "
                    "المزدوجة الحالية.")
    add_bullet(doc, "الدالة query_events(filters, logic, limit) في طبقة التخزين تنفّذ الاستعلام وتعيد "
                    "النتائج المطابقة.")

    add_par(doc, "واجهة برمجة التطبيقات:", bold=True)
    add_code(doc, 'POST /api/forensics/query\n'
                  'البنية: { "filters": [ { "field": "...", "op": "...", "value": ... }, ... ], '
                  '"logic": "AND" | "OR", "limit": 100 }\n'
                  'الاستجابة تتضمن أيضاً available_fields و available_operators لمساعدة الواجهة الأمامية '
                  'على بناء نموذج بحث ديناميكي.')

    add_par(doc, "الصلاحيات (RBAC):", bold=True)
    add_par(doc, "يستخدم نفس صلاحية forensics:read (admin و analyst).")

    # ------------------------------------------------------------------
    # Verification
    # ------------------------------------------------------------------
    add_heading(doc, "التحقق من صحة العمل (Verification)", level=1)

    add_heading(doc, "1. اختبارات آلية (Automated Tests)", level=2)
    add_par(doc, "تمت إضافة ملفين جديدين للاختبارات بنفس أسلوب الاختبارات الحالية في المشروع (pytest):")
    add_bullet(doc, "tests/test_incidents_forensics.py — 12 اختباراً على مستوى طبقة التخزين (دورة حياة "
                    "الحادثة، ربط الأحداث، فلاتر القائمة، إدارة الحالات، صحة سلسلة التدقيق، كشف العبث "
                    "بالسلسلة، ترتيب إعادة التشغيل، النافذة الزمنية، منطق AND/OR في الاستعلام، عامل "
                    "contains، تجاهل الحقول غير المعروفة).")
    add_bullet(doc, "tests/test_incidents_api.py — 6 اختبارات على مستوى واجهات API الكاملة (مع تسجيل دخول "
                    "فعلي وفحص الصلاحيات): دورة حياة الحادثة كاملة عبر API، ربط حادثة بحالة، منع المستخدم "
                    "viewer من الكتابة، سجل التدقيق والتحقق منه، إعادة التشغيل والاستعلام، ومنع viewer من "
                    "الوصول إلى وحدة التحقيق الجنائي.")

    add_par(doc, "نتيجة تنفيذ الاختبارات:", bold=True)
    add_code(doc, "pytest -q\n"
                  "==> 39 passed, 2 warnings\n"
                  "(21 اختباراً سابقاً للنظام + 18 اختباراً جديداً للمهام الخمس)")
    add_par(doc, "ملاحظة: تمت أيضاً إضافة الحزمة المفقودة httpx (المطلوبة لتشغيل FastAPI TestClient) إلى "
                  "requirements.txt، وكانت هذه الحزمة ناقصة حتى بالنسبة لاختبارات النظام السابقة.")

    add_heading(doc, "2. اختبار حي على الخادم الفعلي (Live Smoke Test)", level=2)
    add_par(doc, "تم إعادة تشغيل الخادم (uvicorn على المنفذ 8080) لتحميل المسارات الجديدة، ثم تم تنفيذ "
                  "السيناريو التالي عبر تسجيل دخول admin وجلسة فعلية:")
    add_bullet(doc, "إنشاء حادثة جديدة (POST /api/incidents) ← نجح، وأُرجع كائن الحادثة كاملاً.")
    add_bullet(doc, "تعديل حالة الحادثة وتعيين محقق (PATCH /api/incidents/{id}) ← الحالة تغيّرت بنجاح "
                    "إلى investigating.")
    add_bullet(doc, "إنشاء حالة جديدة وربطها بالحادثة (POST /api/cases) ← نجح، وظهرت الحادثة ضمن incident_ids.")
    add_bullet(doc, "حظر عنوان IP تجريبي (POST /api/block-ip) ← نجح وتم تسجيله في سجل التدقيق تلقائياً.")
    add_bullet(doc, "عرض سجلات التدقيق (GET /api/audit-logs) ← ظهرت 4 سجلات متسلسلة "
                    "(incident_created → incident_updated → case_created → ip_blocked)، كل سجل يحمل "
                    "prev_hash مطابقاً تماماً لـ hash السجل الذي يسبقه.")
    add_bullet(doc, "التحقق من سلامة السلسلة (GET /api/audit-logs/verify) ← النتيجة: "
                    '{ "valid": true, "total": 4, "broken_at": null }')
    add_bullet(doc, "إعادة تشغيل أحداث الجدار الناري (GET /api/forensics/replay?target_entity=firewall) "
                    "← أعاد حدث ip_blocked المسجَّل بترتيب صحيح.")
    add_bullet(doc, "استعلام متقدم (POST /api/forensics/query) بشرط event_type = ip_blocked ← أعاد النتيجة "
                    "الصحيحة مع قوائم available_fields و available_operators.")

    add_par(doc, "النتيجة الإجمالية:", bold=True)
    add_par(doc, "جميع المهام الخمس تعمل بشكل صحيح وكامل، سواء على مستوى الاختبارات الآلية (39/39 ناجحة) "
                  "أو على مستوى الاستخدام الفعلي للخادم الحي، دون أي أعطال أو تراجع في الميزات الموجودة "
                  "مسبقاً في النظام.")

    # ------------------------------------------------------------------
    # Files summary
    # ------------------------------------------------------------------
    add_heading(doc, "ملخص الملفات المعدَّلة والمضافة", level=1)
    add_bullet(doc, "security/roles.py — إضافة 6 صلاحيات جديدة (incidents:read/write، cases:read/write، "
                    "audit:verify، forensics:read).")
    add_bullet(doc, "storage/persistence.py — إضافة 5 جداول جديدة، دالة تجزئة التدقيق، باني الاستعلامات "
                    "الآمن، و~25 دالة جديدة (SQLite).")
    add_bullet(doc, "storage/postgres_store.py — نفس الإضافات بالكامل لقاعدة بيانات PostgreSQL.")
    add_bullet(doc, "backend/routes/incidents.py (ملف جديد) — واجهات API للحوادث والحالات.")
    add_bullet(doc, "backend/routes/forensics.py (ملف جديد) — واجهات API لسجلات التدقيق وإعادة التشغيل "
                    "والاستعلام المتقدم.")
    add_bullet(doc, "backend/main.py — تسجيل المسارات الجديدة (incidents، forensics).")
    add_bullet(doc, "backend/routes/security.py — ربط عمليتي حظر/فك حظر IP بسجل التدقيق.")
    add_bullet(doc, "tests/test_incidents_forensics.py و tests/test_incidents_api.py (ملفان جديدان) — "
                    "30 سيناريو اختبار جديد (12 + 6 اختبارات أساسية موزعة على 5 مهام، بالإضافة لاختبارات "
                    "فرعية).")
    add_bullet(doc, "requirements.txt — إضافة الحزمة الناقصة httpx.")

    add_heading(doc, "الخلاصة", level=1)
    add_par(doc,
            "تم إنجاز المهام الخمس الأولى المسندة إلى مهندس الأنظمة والاستجابة (الفاروق) بالكامل، وفق "
            "البند الثالث من خطة ترقية النظام إلى منصة SOC، مع الالتزام التام ببنية المشروع ونمط الكود "
            "ونظام الصلاحيات الحالي. تم التحقق من صحة جميع الميزات عبر 39 اختباراً آلياً ناجحاً بالكامل "
            "بالإضافة إلى اختبار حي على الخادم الفعلي. النظام جاهز الآن لانتقال بقية الفريق إلى المهام "
            "التالية (4 إلى 9) التي تعتمد على هذه الأساسات (مثل واجهات SOC الأمامية لعرض الحوادث والحالات "
            "وسجلات التدقيق ومحرك إعادة التشغيل).")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
